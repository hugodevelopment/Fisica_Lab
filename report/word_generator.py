"""report/word_generator.py — Gera relatório Word com os resultados."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


AZUL = RGBColor(0x1A, 0x56, 0xA0)


def gerar(
    titulo:       str,
    expressao:    str,
    dados:        dict,
    derivadas:    dict,
    incerteza:    float,
    valor:        float,
    valor_teorico: float = None,
    caminho:      str   = "relatorio.docx",
) -> str:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # Título
    h = doc.add_heading(titulo, level=0)

    p = doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].italic = True

    # 1. Dados
    _secao(doc, "1. Dados Experimentais")
    t = _tabela(doc, ["Variável", "Valor"])
    for sym, val in dados.items():
        _linha(t, [str(sym), f"{val:.6g}"])

    # 2. Expressão e resultado
    _secao(doc, "2. Resultado")
    doc.add_paragraph(f"Expressão:          {expressao}")
    doc.add_paragraph(f"Valor calculado:    {valor:.6g}")
    doc.add_paragraph(f"Incerteza:          {incerteza:.2g}")
    doc.add_paragraph(f"Resultado final:    ({valor:.4g} ± {incerteza:.2g})")

    # 3. Derivadas
    _secao(doc, "3. Derivadas Parciais")
    t2 = _tabela(doc, ["Variável", "∂f/∂xi (simbólica)", "Valor numérico"])
    for var, info in derivadas.items():
        _linha(t2, [str(var), str(info["simbolica"]), f"{info['numerica']:.6g}"])

    # 4. Erro percentual (opcional)
    if valor_teorico is not None:
        erro = abs((valor - valor_teorico) / valor_teorico) * 100
        _secao(doc, "4. Erro Percentual")
        doc.add_paragraph(f"Valor teórico:   {valor_teorico:.6g}")
        p = doc.add_paragraph(f"Erro percentual: {erro:.2f}%")
        if erro > 10:
            p.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.save(str(caminho))
    return caminho


# ── helpers ───────────────────────────────────────────────────────────────────

def _secao(doc, texto):
    doc.add_paragraph()
    h = doc.add_heading(texto, level=1)


def _tabela(doc, cabecalho):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    for i, c in enumerate(cabecalho):
        cell = t.rows[0].cells[i]
        cell.text = c
        cell.paragraphs[0].runs[0].bold = True
    return t


def _linha(tabela, valores):
    linha = tabela.add_row().cells
    for i, v in enumerate(valores):
        linha[i].text = v
